#!/usr/bin/env python3
"""Build the polished English M1 supervisor update from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "202A35"
MUTED = "66717D"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "EDF4FA"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_fixed_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for idx, cell in enumerate(row.cells):
            width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_runs(paragraph, text: str, *, default_bold=False, color=INK) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    for part in token_re.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, color=color, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, color=color, bold=default_bold)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.92)
    section.right_margin = Inches(0.92)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, before, after, color in (
        ("Heading 1", 15, 14, 6, BLUE),
        ("Heading 2", 12.5, 10, 5, BLUE),
        ("Heading 3", 11.5, 8, 4, DARK_BLUE),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.10

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("TRANSFER VS. RELEARNING  |  M1 COMPLETION UPDATE")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    label = fp.add_run("13 July 2026   |   ")
    set_run_font(label, size=9, color=MUTED)
    add_page_field(fp)


def add_masthead(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("M1 COMPLETION UPDATE")
    set_run_font(r, size=23, color=INK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("Validated capacity result for English factual acquisition")
    set_run_font(r, size=13, color=MUTED)

    metadata = [
        ("To", "Max / Thesis Supervisor"),
        ("Date", "13 July 2026"),
        ("Project", "Transfer vs. Relearning in Cross-Lingual Factual Adaptation"),
        ("Status", "M1 gate passed; ready to proceed to M2/M3"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{label}: ")
        set_run_font(r, size=10.5, color=INK, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=10.5, color=INK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(8)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_box_paragraph(paragraph, fill: str, border: str = "808080") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "5")
        node.set(qn("w:color"), border)
        p_bdr.append(node)
    p_pr.append(p_bdr)


def add_callout(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(7)
    p.paragraph_format.right_indent = Pt(7)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.12
    style_box_paragraph(p, CALLOUT)
    add_inline_runs(p, text, default_bold=True, color=DARK_BLUE)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(7)
    p.paragraph_format.right_indent = Pt(7)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.0
    style_box_paragraph(p, LIGHT_GRAY)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, name="Consolas", size=9.2, color=INK)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    widths = [TABLE_WIDTH // cols] * cols
    remainder = TABLE_WIDTH - sum(widths)
    widths[-1] += remainder
    if cols >= 3:
        first = 2500
        rest = (TABLE_WIDTH - first) // (cols - 1)
        widths = [first] + [rest] * (cols - 1)
        widths[-1] += TABLE_WIDTH - sum(widths)

    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_fixed_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        for cidx, value in enumerate(row):
            cell = table.cell(ridx, cidx)
            if ridx == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if ridx < len(rows) - 1:
                p.paragraph_format.keep_with_next = True
            if cidx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_inline_runs(p, value, default_bold=(ridx == 0), color=INK)
            for run in p.runs:
                run.font.size = Pt(9.2 if cols >= 5 else 9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def new_numbering_instance(doc: Document) -> int:
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    numbering = doc.part.numbering_part.element
    source_num = next(num for num in numbering.num_lst if num.numId == style_num_id)
    abstract_id = source_num.abstractNumId.val
    new_num = numbering.add_num(abstract_id)
    new_num.add_lvlOverride(ilvl=0).add_startOverride(1)
    return new_num.numId


def apply_numbering(paragraph, num_id: int) -> None:
    num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        parts = [part.strip() for part in lines[idx].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            rows.append(parts)
        idx += 1
    return rows, idx


def gather_paragraph(lines: list[str], start: int) -> tuple[str, int]:
    chunks = []
    idx = start
    while idx < len(lines):
        s = lines[idx].strip()
        if not s or s.startswith(("#", "- ", ">", "```", "|")) or re.match(r"^\d+\.\s", s):
            break
        chunks.append(s)
        idx += 1
    return " ".join(chunks), idx


def add_body_from_markdown(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    skipped_source_title = False
    active_numbering_id = None
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()
        if not stripped:
            active_numbering_id = None
            idx += 1
            continue
        if stripped.startswith("# ") and not skipped_source_title:
            skipped_source_title = True
            idx += 1
            while idx < len(lines) and (not lines[idx].strip() or lines[idx].strip().startswith("**")):
                idx += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline_runs(p, stripped[3:], default_bold=True, color=BLUE)
            idx += 1
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 2")
            add_inline_runs(p, stripped[4:], default_bold=True, color=BLUE)
            idx += 1
            continue
        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue
        if stripped.startswith(">"):
            quote = []
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote.append(lines[idx].strip()[1:].strip())
                idx += 1
            add_callout(doc, " ".join(quote))
            continue
        if stripped.startswith("- "):
            text = stripped[2:]
            idx += 1
            while idx < len(lines) and lines[idx].startswith("  ") and not lines[idx].strip().startswith("-"):
                text += " " + lines[idx].strip()
                idx += 1
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            continue
        numbered = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if numbered:
            text = numbered.group(2)
            idx += 1
            while idx < len(lines) and lines[idx].startswith("   "):
                text += " " + lines[idx].strip()
                idx += 1
            p = doc.add_paragraph(style="List Number")
            if active_numbering_id is None:
                active_numbering_id = new_numbering_instance(doc)
            apply_numbering(p, active_numbering_id)
            add_inline_runs(p, text)
            continue

        text, next_idx = gather_paragraph(lines, idx)
        if not text:
            text = stripped
            next_idx = idx + 1
        p = doc.add_paragraph()
        add_inline_runs(p, text)
        idx = next_idx


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_supervisor_m1_docx.py INPUT.md OUTPUT.docx")
    source = Path(sys.argv[1])
    output = Path(sys.argv[2])
    doc = Document()
    configure_document(doc)
    add_masthead(doc)
    add_body_from_markdown(doc, source.read_text(encoding="utf-8"))

    props = doc.core_properties
    props.title = "M1 Completion Update"
    props.subject = "Supervisor briefing on English factual acquisition completion"
    props.author = "Umut Yunus Yesildal"
    props.keywords = "thesis, M1, factual acquisition, SmolLM2, cross-lingual transfer"
    props.comments = "Prepared for thesis supervision discussion."

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
